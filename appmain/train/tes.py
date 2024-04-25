#
#
# from docx import Document
#
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from bs4 import BeautifulSoup
# import base64
#
#
# def convert_docx_to_html(docx_file_path):
#     # Открываем файл .docx
#     doc = Document(docx_file_path)
#
#     # Создаем объект BeautifulSoup для хранения HTML
#     html_content = BeautifulSoup(features="html.parser")
#
#     # Создаем контейнер для содержимого документа
#     body = html_content.new_tag('body')
#     html_content.append(body)
#
#     # Проходим по всем параграфам в документе
#     for paragraph in doc.paragraphs:
#         # Создаем HTML-тег для каждого параграфа
#         p_tag = html_content.new_tag('p')
#
#         # Устанавливаем выравнивание текста
#         if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
#             p_tag['style'] = 'text-align: center;'
#         elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
#             p_tag['style'] = 'text-align: right;'
#
#         # Добавляем текст параграфа в HTML-тег
#         p_tag.string = paragraph.text
#
#         # Добавляем HTML-тег в контейнер body
#         body.append(p_tag)
#
#     # Проходим по всем изображениям в документе
#     for image in doc.inline_shapes:
#         # Создаем HTML-тег для изображения
#         img_tag = html_content.new_tag('img')
#
#         # Получаем бинарные данные изображения
#         image_data = image._inline.graphic.graphicData.pic.blipFill.blip.embed
#         image_stream = doc.part.related_parts[image_data].blob
#         image_base64 = base64.b64encode(image_stream).decode('utf-8')
#
#         # Устанавливаем атрибуты src и alt для тега изображения
#         img_tag['src'] = f'data:image/png;base64,{image_base64}'
#         img_tag['alt'] = 'Image'
#
#         # Добавляем HTML-тег изображения в контейнер body
#         body.append(img_tag)
#
#     # Возвращаем HTML-контент в виде строки
#     return str(html_content)
#
# # Путь к вашему файлу .docx
# docx_file_path = r'D:\КОЛЯМБА\НикитаДиплом\кинематика\теория\кинематика теория.docx'
#
# # Преобразуем содержимое .docx в HTML
# html_content = convert_docx_to_html(docx_file_path)
#
# # Выводим HTML-контент
# print(html_content)


from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import base64

def convert_docx_to_html(docx_file_path):
    # Открываем файл .docx
    doc = Document(docx_file_path)

    # Создаем объект BeautifulSoup для хранения HTML
    html_content = BeautifulSoup(features="html.parser")

    # Создаем контейнер для содержимого документа
    body = html_content.new_tag('body')
    html_content.append(body)

    # Список для хранения изображений
    images = []

    # Проходим по всем параграфам в документе
    for paragraph in doc.paragraphs:
        # Создаем HTML-тег для каждого параграфа
        p_tag = html_content.new_tag('p')

        # Устанавливаем выравнивание текста
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            p_tag['style'] = 'text-align: center;'
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            p_tag['style'] = 'text-align: right;'

        # Добавляем текст параграфа в HTML-тег
        p_tag.string = paragraph.text

        # Добавляем HTML-тег в контейнер body
        body.append(p_tag)

    # Проходим по всем изображениям в документе и добавляем их в список в том порядке, в котором они встречаются
    for image in doc.inline_shapes:
        # Получаем бинарные данные изображения
        image_data = image._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_stream = doc.part.related_parts[image_data].blob
        image_base64 = base64.b64encode(image_stream).decode('utf-8')

        # Добавляем данные изображения в список
        images.append(image_base64)

    # Добавляем изображения в HTML в том порядке, в котором они встречаются в списке
    for image_base64 in images:
        # Создаем HTML-тег для изображения
        img_tag = html_content.new_tag('img')

        # Устанавливаем атрибуты src и alt для тега изображения
        img_tag['src'] = f'data:image/png;base64,{image_base64}'
        img_tag['alt'] = 'Image'

        # Добавляем HTML-тег изображения в контейнер body
        body.append(img_tag)

    # Возвращаем HTML-контент в виде строки
    return str(html_content)

# Путь к вашему файлу .docx
docx_file_path = r'D:\КОЛЯМБА\НикитаДиплом\кинематика\теория\кинематика теория.docx'

# Преобразуем содержимое .docx в HTML
html_content = convert_docx_to_html(docx_file_path)

# Выводим HTML-контент
print(html_content)